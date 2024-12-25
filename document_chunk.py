import os 
import json 
from threading import Thread 
from unittest import result 
from flask import Blueprint, request, jsonify 
from langchain.text_splitter import RecursiveCharacterTextSplitter 
from PyPDF2 import PdfReader 
from langchain.schema import Document 
import embedding_from_chunk as embed 

chunk_page = Blueprint('chunk_page', __name__) 
UPLOAD_FOLDER = 'uploaded_documents' 
CHUNK_FOLDER = 'chunked_documents' 
chunking_status = {}  # To track the status of each file 


def is_file_chunked(filename): 
    chunked_file_path = os.path.join(CHUNK_FOLDER, f"{filename}.json") 
    return os.path.exists(chunked_file_path) 

def extract_text_from_pdf(file_path): 
    try: 
        reader = PdfReader(file_path) 
        text = [] 
        for i,page in enumerate(reader.pages): 
            text.append({"page_content":page.extract_text(),"page_number": i+1,"chunk_id": i}) 
        return text 
    except Exception as e: 
        raise RuntimeError(f"Error extracting text from PDF: {str(e)}") 
 
def extract_text_from_docx(file_path): 
    try: 
        from docx import Document as DocxDocument 
        doc = DocxDocument(file_path) 
        #doc = Document(file_path) 
        text = []     
        for i,paragraph in enumerate(doc.paragraphs): 
            text.append({"page_content":paragraph.text,"page_number":i+1,"chunk_id": i}) 
        return text 
    except Exception as e: 
        raise RuntimeError(f"Error extracting text from DOCX: {str(e)}") 
 
def extract_text_from_general(file_path): 
    with open(file_path,'r',encoding='utf-8') as f: 
        text = [{"page_content":f.read(),"page_number":1,"chunk_id": 1}]
     
    return text 

def chunk_file(file_path, output_file,storage_type): 
    # Determine file type based on extension 
    _, file_extension = os.path.splitext(file_path) 
    file_extension = file_extension.lower() 
    file_name= os.path.basename(file_path)   

    if file_extension == ".pdf": 
        text = extract_text_from_pdf(file_path) 
     
    elif file_extension == ".docx": 
        text = extract_text_from_docx(file_path) 
    elif file_extension == ".txt": 
      
        text = extract_text_from_general(file_path) 
    else: 
        raise ValueError(f"Unsupported file type: {file_extension}") 
    
    # Use RecursiveCharacterTextSplitter to split the text 
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200) 
    chunks = [] 
    print("document",text)
    for item in text: 
     
        document = Document(page_content=item["page_content"]) 
     
        split_chunks = text_splitter.split_documents([document]) 
        for chunk in split_chunks: 
            chunks.append({ 
                "page_content":chunk.page_content, 
                "page_number":item["page_number"], 
                "chunk_id":item["chunk_id"],
                "filename":file_name 
            }) 
             
  
    try: 
    # Save chunks to JSON 
        with open(output_file, 'w', encoding='utf-8') as f: 
            print("error code") 
            json.dump(chunks, f, ensure_ascii=False, indent=4) 
            print("error code end") 
    except Exception as e: 
         raise ValueError(f"error serialize  chunks : {e}") 
    chunking_status[file_name] = "Embedding in Progress" 
    # start Embedding process 
    embed.embed_chunk_from_file(output_file,storage_type) 
    chunking_status[file_name] = "Embedding in completed" 
 
    
    #return chunks 

def process_files_in_background(filenames,storage_type): 
    for filename in filenames: 
        file_path = os.path.join(UPLOAD_FOLDER, filename) 
        os.makedirs(CHUNK_FOLDER, exist_ok=True) 
        chunked_file_path = os.path.join(CHUNK_FOLDER, f"{filename}.json") 
        chunking_status[filename] = "In Process" 
        try: 
            chunk_file(file_path, chunked_file_path,storage_type) 
            chunking_status[filename] = "Chunking Complete" 
        except Exception as e: 
            chunking_status[filename] = f"Error: {str(e)}" 

@chunk_page.route('/process', methods=['POST']) 
def process_chunk(): 
    filenames = request.json.get('filenames') 
    storage_type = request.json.get('storage_type',"faiss")   
    if not filenames: 
        return jsonify({"message": "No filenames provided"}), 400 
  
    for filename in filenames: 
        file_path =os.path.join(UPLOAD_FOLDER,filename) 
        chunk_file_path = os.path.join(CHUNK_FOLDER,file_path,f"{filename}.json") 
        if not  os.path.exists(file_path): 
            result[filename]="file not found" 
            continue 
     # Start background processing 
    print ("filename",filenames)
    process_files_in_background(filenames,storage_type) 
    # Initialize status for files 
    for filename in filenames: 
        chunking_status[filename] = "Chunking and embedding done " 
    return jsonify({"message": "Chunking started in background", "status": chunking_status[filename]})  

@chunk_page.route('/status', methods=['GET'])



def get_status():
    """
    Returns the current status of chunking as JSON.
    """
    try:
        # Ensure data is JSON serializable
        serializable_status = {
            key: str(value) if callable(value) else value
            for key, value in chunking_status.items()
        }

        # Return status as a JSON response
        return jsonify(serializable_status), 200
    except Exception as e:
        print(f"Error in get_status: {e}")
        return jsonify({"error": "Internal Server Error", "message": str(e)}), 500
    